#!/usr/bin/env python3
"""
Simple script to convert the OCI deployment guide to HTML format
"""

import markdown2
import os
from pathlib import Path

def create_formatted_html(md_content):
    """Convert markdown to properly formatted HTML"""
    # Convert markdown to HTML
    html_body = markdown2.markdown(md_content, extras=['tables', 'fenced-code-blocks', 'header-ids'])
    
    # Create full HTML document with styling
    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>QuickShop E-commerce Platform - Complete Deployment Guide</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            line-height: 1.6;
            margin: 0;
            padding: 20px;
            color: #333;
            max-width: 1200px;
            margin: 0 auto;
            background-color: #fff;
        }}
        .header {{
            text-align: center;
            border-bottom: 3px solid #2c3e50;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }}
        h1 {{
            color: #2c3e50;
            font-size: 2.5em;
            margin-bottom: 10px;
        }}
        h2 {{
            color: #3498db;
            border-bottom: 1px solid #eee;
            padding-bottom: 10px;
            margin-top: 30px;
        }}
        h3 {{
            color: #2c3e50;
            margin-top: 25px;
        }}
        h4 {{
            color: #7f8c8d;
        }}
        code {{
            background-color: #f8f9fa;
            padding: 2px 6px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
            color: #e74c3c;
            border: 1px solid #e9ecef;
        }}
        pre {{
            background-color: #2c3e50;
            color: #f8f9fa;
            padding: 15px;
            border-radius: 5px;
            overflow-x: auto;
            font-family: 'Courier New', monospace;
            border-left: 4px solid #3498db;
        }}
        pre code {{
            background: none;
            color: inherit;
            padding: 0;
            border: none;
        }}
        a {{
            color: #3498db;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }}
        th {{
            background-color: #3498db;
            color: white;
        }}
        tr:nth-child(even) {{
            background-color: #f2f2f2;
        }}
        ul, ol {{
            padding-left: 20px;
        }}
        li {{
            margin: 8px 0;
        }}
        .toc-container {{
            background-color: #f8f9fa;
            border-left: 4px solid #3498db;
            padding: 15px;
            margin: 20px 0;
        }}
        .footer {{
            margin-top: 40px;
            text-align: center;
            color: #7f8c8d;
            font-size: 0.9em;
            border-top: 1px solid #eee;
            padding-top: 20px;
        }}
        .highlight {{
            background-color: #fff3cd;
            border-left: 4px solid #ffc107;
            padding: 15px;
            margin: 20px 0;
            border-radius: 3px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>QuickShop E-commerce Platform</h1>
        <h2>Complete Deployment Guide</h2>
        <p>Comprehensive guide for deploying QuickShop on Oracle Cloud Infrastructure</p>
    </div>
    {html_body}
    <div class="footer">
        <p>QuickShop Deployment Guide - Version 1.0.0 - December 2024</p>
        <p>© QuickShop Team. All rights reserved.</p>
    </div>
</body>
</html>"""
    
    return html_doc

def main():
    # Read the markdown file
    md_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.md"

    if not os.path.exists(md_file_path):
        print(f"Markdown file not found: {md_file_path}")
        return

    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Convert to HTML
    html_content = create_formatted_html(md_content)

    # Save HTML file
    html_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.html"
    with open(html_file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

    print(f"HTML document created: {html_file_path}")

    # Create a basic Word document as text with some formatting
    docx_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.docx"

    try:
        from docx import Document
        from docx.shared import Inches

        # Create a new document
        doc = Document()

        # Add title
        doc.add_heading('QuickShop E-commerce Platform', 0)
        doc.add_heading('Complete Deployment Guide', 1)

        # Add a paragraph with description
        doc.add_paragraph('Comprehensive guide for deploying QuickShop on Oracle Cloud Infrastructure')

        # Add page break
        doc.add_page_break()

        # Process markdown content - split by lines and add to document
        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=4)
            elif line.startswith('- '):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith('1. '):
                doc.add_paragraph(line[3:], style='List Number')
            elif line.startswith('```'):
                continue  # Skip code block markers
            elif line == '':
                continue  # Skip empty lines
            elif line.startswith('   ') or line.startswith('```'):
                # Add code blocks
                doc.add_paragraph(line.strip(), style='Code')
            else:
                doc.add_paragraph(line)

        # Save document
        doc.save(docx_file_path)
        print(f"Word document created: {docx_file_path}")
    except ImportError:
        print("python-docx not available, skipping Word document creation")
        docx_file_path = None

    # Also save as plain text with better formatting
    txt_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.txt"
    with open(txt_file_path, 'w', encoding='utf-8') as f:
        f.write("="*80 + "\n")
        f.write("QuickShop E-commerce Platform - Complete Deployment Guide\n")
        f.write("="*80 + "\n\n")
        f.write(md_content)
        f.write("\n\n" + "="*80)
        f.write("\nQuickShop Deployment Guide - Version 1.0.0 - December 2024")
        f.write("\n" + "="*80)

    print(f"Text document created: {txt_file_path}")

    print("\nTo convert HTML to PDF:")
    print("- Open the HTML file in a web browser")
    print("- Print the page (Cmd+P on Mac, Ctrl+P on Windows)")
    print("- Select 'Save as PDF' as the destination")

if __name__ == "__main__":
    main()