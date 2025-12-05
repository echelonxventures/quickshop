#!/usr/bin/env python3
"""
Script to convert the OCI deployment guide to PDF and DOCX formats
"""

import markdown
import pdfkit
import os
from docx import Document
from docx.shared import Inches
from pathlib import Path

def markdown_to_html(md_content):
    """Convert markdown to HTML with proper formatting"""
    # Custom conversion to handle our specific markdown
    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>QuickShop OCI Deployment Guide</title>
    <style>
        body {{
            font-family: Arial, sans-serif;
            line-height: 1.6;
            margin: 40px;
            background-color: white;
        }}
        h1, h2, h3, h4 {{
            color: #2c3e50;
        }}
        h1 {{
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        code {{
            background-color: #f4f4f4;
            padding: 2px 4px;
            border-radius: 3px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background-color: #f4f4f4;
            padding: 10px;
            border-radius: 5px;
            overflow-x: auto;
            border-left: 4px solid #3498db;
        }}
        a {{
            color: #3498db;
            text-decoration: none;
        }}
        a:hover {{
            text-decoration: underline;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 8px;
            text-align: left;
        }}
        th {{
            background-color: #f2f2f2;
        }}
    </style>
</head>
<body>
{html}
</body>
</html>"""

def markdown_to_docx(md_content, docx_path):
    """Convert markdown content to DOCX format"""
    doc = Document()
    
    # Add title page
    doc.add_heading('QuickShop E-commerce Platform - Complete Deployment Guide', 0)
    doc.add_paragraph('Complete guide for deploying QuickShop on Oracle Cloud Infrastructure')
    doc.add_paragraph()
    
    # Simple conversion - split by lines and handle basic formatting
    lines = md_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('```'):
            # Skip code block markers, add the actual code in next lines
            continue
        elif line.startswith('1. '):
            doc.add_paragraph(line[3:], style='List Number')
        elif line == '':
            continue
        else:
            # Just add as paragraph
            if not any(line.startswith(tag) for tag in ['#', '##', '###', '####', '- ', '1.']):
                doc.add_paragraph(line)
    
    doc.save(docx_path)
    print(f"DOCX document saved to: {docx_path}")

def main():
    # Read the markdown file
    md_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.md"
    
    if not os.path.exists(md_file_path):
        print(f"Markdown file not found: {md_file_path}")
        return
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert to HTML first
    html_content = markdown_to_html(md_content)
    html_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.html"
    
    with open(html_file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)
    
    # Convert HTML to PDF
    pdf_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.pdf"
    
    try:
        # Try to create PDF using weasyprint if available
        try:
            from weasyprint import HTML, CSS
            HTML(string=html_content).write_pdf(pdf_file_path)
            print(f"PDF document created using WeasyPrint: {pdf_file_path}")
        except ImportError:
            # Fallback to pdfkit if wkhtmltopdf is available
            try:
                options = {
                    'page-size': 'A4',
                    'margin-top': '0.75in',
                    'margin-right': '0.75in',
                    'margin-bottom': '0.75in',
                    'margin-left': '0.75in',
                    'encoding': "UTF-8",
                    'no-outline': None
                }
                pdfkit.from_string(html_content, pdf_file_path, options=options)
                print(f"PDF document created using pdfkit: {pdf_file_path}")
            except OSError:
                print("PDF creation failed - neither WeasyPrint nor wkhtmltopdf available")
                print("Creating PDF using alternative method...")
                # Create a simple text version
                with open(pdf_file_path.replace('.pdf', '_fallback.txt'), 'w', encoding='utf-8') as f:
                    f.write(md_content)
                print(f"Fallback text file created: {pdf_file_path.replace('.pdf', '_fallback.txt')}")
        
        # Create DOCX
        docx_file_path = "/Users/bipin/Desktop/quickshop/OCI_Deployment_Guide.docx"
        markdown_to_docx(md_content, docx_file_path)
        
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        # Create fallback files
        fallback_pdf = pdf_file_path.replace('.pdf', '_fallback.txt')
        with open(fallback_pdf, 'w', encoding='utf-8') as f:
            f.write(md_content)
        print(f"Fallback file created: {fallback_pdf}")

if __name__ == "__main__":
    main()